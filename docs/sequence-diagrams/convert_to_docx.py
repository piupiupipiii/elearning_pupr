from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Get script directory
script_dir = os.path.dirname(os.path.abspath(__file__))

# Create document
doc = Document()

# Title
title = doc.add_heading('Sequence Diagram System E-Learning SMKK', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

intro = doc.add_paragraph('Dokumen ini berisi 6 sequence diagram sesuai standar UML untuk laporan Kerja Praktek.')
intro.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 6 Sequence diagrams sesuai aturan UML
diagrams = [
    {
        'title': '1. Sequence Diagram: Melakukan Autentikasi Pengguna',
        'desc': 'Proses login dengan validasi kredensial, pembentukan session, dan penanganan kondisi berhasil/gagal.',
        'image': '01_autentikasi.png',
        'steps': [
            'Pengguna mengakses halaman login',
            'Sistem menampilkan form login',
            'Pengguna mengisi email dan password, klik Masuk',
            'AuthController memvalidasi kredensial',
            'Jika valid: session dibuat dan redirect ke Pengantar Modul',
            'Jika tidak valid: tampilkan pesan error'
        ]
    },
    {
        'title': '2. Sequence Diagram: Mengakses Pengantar Modul',
        'desc': 'Proses akses halaman utama dengan pengecekan status autentikasi untuk menentukan tombol Login atau Mulai.',
        'image': '02_pengantar_modul.png',
        'steps': [
            'Pengguna mengakses halaman utama',
            'Controller mengecek status autentikasi',
            'Jika sudah login: tampilkan tombol "Mulai"',
            'Jika belum login: tampilkan tombol "Login"'
        ]
    },
    {
        'title': '3. Sequence Diagram: Mengakses Sub Menu Materi',
        'desc': 'Proses menampilkan daftar materi dengan status progress (terkunci, terbuka, selesai) untuk setiap user.',
        'image': '03_submenu_materi.png',
        'steps': [
            'Pengguna mengakses sub menu materi',
            'MaterialController mengambil daftar materi',
            'Controller mengambil data progress user',
            'Jika user baru: inisialisasi progress, materi pertama dibuka',
            'Controller menentukan status tiap materi',
            'Tampilkan slider materi dengan status'
        ]
    },
    {
        'title': '4. Sequence Diagram: Mengakses dan Menonton Video Materi',
        'desc': 'Proses pemutaran video materi termasuk validasi akses dan deteksi video selesai.',
        'image': '04_menonton_video.png',
        'steps': [
            'Pengguna memilih materi dari slider',
            'MaterialController mengecek akses user',
            'Jika terkunci: tampilkan pesan akses ditolak',
            'Jika terbuka: tampilkan video player',
            'User menonton video hingga selesai',
            'Sistem mengaktifkan tombol Next ke Quiz'
        ]
    },
    {
        'title': '5. Sequence Diagram: Mengerjakan Kuis Evaluasi',
        'desc': 'Proses quiz dengan pengambilan soal acak, validasi jawaban, perhitungan skor, dan unlock materi berikutnya.',
        'image': '05_kuis_evaluasi.png',
        'steps': [
            'Pengguna mengakses halaman quiz',
            'QuizController mengambil 2 soal acak',
            'User menjawab soal, sistem validasi dan tampilkan feedback',
            'User submit quiz',
            'Controller menyimpan jawaban dan menghitung skor',
            'Update progress menjadi selesai',
            'Unlock materi berikutnya',
            'Redirect ke sub menu'
        ]
    },
    {
        'title': '6. Sequence Diagram: Mengelola Profil Pengguna',
        'desc': 'Proses menampilkan profil dengan agregasi data progress pembelajaran dan riwayat nilai quiz.',
        'image': '06_profil_pengguna.png',
        'steps': [
            'Pengguna mengakses halaman profil',
            'ProfileController mengambil jumlah materi selesai',
            'Controller mengambil total materi',
            'Controller mengambil riwayat nilai quiz',
            'Controller menghitung rata-rata nilai',
            'Tampilkan profil dengan progress dan riwayat nilai'
        ]
    }
]

for d in diagrams:
    # Title
    doc.add_heading(d['title'], level=1)
    
    # Description
    doc.add_paragraph(d['desc'])
    
    # Image
    img_path = os.path.join(script_dir, d['image'])
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(6))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Steps
    doc.add_heading('Langkah-langkah:', level=2)
    for i, step in enumerate(d['steps'], 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_paragraph()  # Spacing

# Summary table
doc.add_heading('Ringkasan Sequence Diagram', level=1)
table = doc.add_table(rows=7, cols=4)
table.style = 'Table Grid'

# Header row
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'No'
hdr_cells[1].text = 'Sequence Diagram'
hdr_cells[2].text = 'Controller'
hdr_cells[3].text = 'Fungsi Utama'

# Data rows
data = [
    ('1', 'Melakukan Autentikasi', 'AuthController', 'Validasi login, buat session'),
    ('2', 'Mengakses Pengantar Modul', 'Controller', 'Cek status autentikasi'),
    ('3', 'Mengakses Sub Menu Materi', 'MaterialController', 'Ambil materi & progress'),
    ('4', 'Menonton Video Materi', 'MaterialController', 'Validasi akses, putar video'),
    ('5', 'Mengerjakan Kuis Evaluasi', 'QuizController', 'Validasi jawaban, hitung skor'),
    ('6', 'Mengelola Profil Pengguna', 'ProfileController', 'Agregasi progress & nilai'),
]

for i, row_data in enumerate(data, 1):
    row_cells = table.rows[i].cells
    for j, text in enumerate(row_data):
        row_cells[j].text = text

doc.add_paragraph()
doc.add_paragraph('Total: 6 Sequence Diagram')

# Save
output_path = os.path.join(script_dir, 'sequence_diagram_smkk.docx')
doc.save(output_path)
print(f'Document saved to {output_path}')
